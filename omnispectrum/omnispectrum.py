import os
import sys
import argparse
import numpy as np
import matplotlib.pyplot as plt
import scipy
import scipy.linalg as la
import pandas as pd
import time
import multiprocessing
from tqdm import tqdm
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import shared

class spectrum:

    def __init__(self, acc:list, t:float, baseline_correction: bool = False):  # 路径
        self.acc = np.array([])  # 地震波a
        self.vel = np.array([])  # 地震波速度
        self.dis = np.array([])  # 地震波位移
        self.tf = np.array([])  # 调幅后地震波a

        self.u = np.array([])
        self.u0 = np.array([])
        self.u00 = np.array([])

        # 真最大最小值用于区分方向时使用
        self.saT = np.array([])
        self.sa = np.array([])  # 谱加速度
        self.sa_truemax = np.array([])  # 真最大值
        self.sa_truemin = np.array([])  # 真最小值

        self.sv = np.array([])  # 谱速度
        self.sv_truemax = np.array([])  # 真最大值
        self.sv_truemin = np.array([])  # 真最小值

        self.sd = np.array([])  # 谱位移
        self.sd_truemax = np.array([])  # 真最大值
        self.sd_truemin = np.array([])  # 真最小值

        self.PGA = 0
        self.PGD = 0
        self.PGV = 0
        self.t = 0
        self.id = ""
        self.x = np.array([])  # 时程

        self.t = t
        self.acc = np.array(acc)
        self.x = np.linspace(0, self.t * len(self.acc), len(self.acc))
        if baseline_correction:
            # 基线修正
            self.sw = self.acc
            self.x = self.x
            A = np.vstack([self.x**0, self.x**1, self.x**2])
            sol, r, rank, s = la.lstsq(A.T, self.acc)
            y_fit = sol[0] + sol[1] * self.x + sol[2] * self.x**2
            self.acc = self.acc - y_fit

        return

    def newmark_beta(self, T, ksi=0.05, gamma=0.5, beta=0.25):  # 步长，周期，地震波
        if T == 0:
            T = 0.00000001
        n = len(self.acc)
        m = 1.0
        c = 2 * m * ksi * (2 * np.pi / T)
        # 分别为加速度，速度，位移
        u00 = np.zeros(n)
        u0 = np.zeros(n)
        u = np.zeros(n)
        k = (2 * np.pi / T) ** 2 * m
        u00[0] = -self.acc[0] - c * u0[0] - k * u[0]
        a1 = m / (beta * (self.t ** 2)) + gamma * c / (beta * self.t)
        a2 = m / (beta * self.t) + (gamma / beta - 1) * c
        a3 = (1 / (2 * beta) - 1) * m + self.t * (gamma / (2 * beta) - 1) * c
        k_hat = k + a1  # k^/hat

        for i in range(1, n):
            p_hat = -self.acc[i] + a1 * u[i - 1] + a2 * u0[i - 1] + a3 * u00[i - 1]
            u[i] = p_hat / k_hat
            u0[i] = gamma / (beta * self.t) * (u[i] - u[i - 1]) + (1 - gamma / beta) * u0[i - 1] + self.t * (
                        1 - gamma / (2 * beta)) * u00[i - 1]
            u00[i] = 1 / (beta * (self.t ** 2)) * (u[i] - u[i - 1]) - u0[i - 1] / (beta * self.t) - (
                        1 / (2 * beta) - 1) * u00[i - 1]
        self.u = u
        self.u0 = u0
        self.u00 = u00 + self.acc
        return u, u0, u00 + self.acc

    def central_difference(self, T, ksi=0.05, gamma=0.5, beta=0.25):
        m = 1.0
        n = len(self.acc)
        u00 = np.zeros(n)
        u0 = np.zeros(n)
        u = np.zeros(n + 1)
        c = 2 * m * ksi * (2 * np.pi / T)
        k = (2 * np.pi / T) ** 2 * m
        u00[0] = -self.acc[0] - c * u0[0] - k * u[0]
        u_1 = u[0] - self.t * u0[0] + self.t ** 2 * u00[0] / 2
        k_hat = m / (self.t ** 2) + c / (2 * self.t)
        a = m / (self.t ** 2) - c / (2 * self.t)
        b = k - 2 * m / (self.t ** 2)
        for i in range(n):
            if i == 1:
                p_hat = -self.acc[i] - a * u_1 - b * u[i]
            else:
                p_hat = -self.acc[i] - a * u[i - 1] - b * u[i]
            u[i + 1] = p_hat / k_hat
            u0[i] = (u[i + 1] - u[i - 1]) / (2 * self.t)
            u00[i] = (u[i + 1] - 2 * u[i] + u[i - 1]) / (self.t ** 2)
        self.u = u[:n]
        self.u0 = u0
        self.u00 = u00 + self.acc
        return u, u0, u00 + self.acc

    def get_PGA(self):
        self.PGA = max(abs(self.acc))
        return max(abs(self.acc))

    def get_PGV(self):
        self.get_v()
        self.PGV = max(abs(self.vel))
        return max(abs(self.vel))

    def get_PGD(self):
        self.get_u()
        self.PGD = max(abs(self.dis))
        return max(abs(self.dis))

    def get_sa(self, begin, end, step):
        T = np.arange(begin, end, step)
        sa = np.array([])
        sv = np.array([])
        su = np.array([])
        sa_truemax = np.array([])
        sa_truemin = np.array([])
        sv_truemax = np.array([])
        sv_truemin = np.array([])
        su_truemax = np.array([])
        su_truemin = np.array([])
        for i in T:
            u, v, a = self.newmark_beta(i)
            sa = np.append(sa, max(abs(a)))
            sa_truemax = np.append(sa_truemax, max(a))
            sa_truemin = np.append(sa_truemin, min(a))

            sv = np.append(sv, max(abs(v)))
            sv_truemax = np.append(sv_truemax, max(v))
            sv_truemin = np.append(sv_truemin, min(v))

            su = np.append(su, max(abs(u)))
            su_truemax = np.append(su_truemax, max(u))
            su_truemin = np.append(su_truemin, min(u))
        self.saT = T
        self.sa = sa
        self.sd = su
        self.sv = sv
        self.sa_truemax = sa_truemax
        self.sa_truemin = sa_truemin
        self.sv_truemax = sv_truemax
        self.sv_truemin = sv_truemin
        self.sd_truemax = su_truemax
        self.sd_truemin = su_truemin
        return

    def get_v(self):  # 步长，地震波
        if len(self.vel) == 0: 
            v = []
            v.append(0)
            for i in range(len(self.acc) - 1):
                v.append(v[i] + self.t * 0.5 * (self.acc[i] + self.acc[i + 1]))
            self.vel = np.array(v)
        return self.vel

    def get_u(self):
        if len(self.dis) == 0:
            v = self.get_v()
            u = []
            u.append(0)
            for i in range(len(self.acc) - 1):
                u.append(u[i] + self.t * v[i] + (0.5 ** 2) * (self.acc[i] + self.acc[i + 1]) * (self.t ** 2))
            self.dis = np.array(u)
        return self.dis

    def tiaofu_sa(self, T, target):  # 地震波步长 地震波 要调幅到的周期和值'
        u, v, a = self.newmark_beta(T)
        sa = max(abs(a))
        self.tf = self.acc
        self.acc = self.acc * (target / sa)
        self.get_v()
        self.get_u()
        return target / sa

    def four(self, form=1):  # 加速度or速度
        if form == 1:
            tmp = self.acc
        elif form == 2:
            tmp = self.vel
        else:
            return
        inpf = np.fft.fft(tmp)  # 默认的是对行向量fourier变换，所以此处要定义下轴
        fs = 1 / self.t
        f = fs * np.arange(0, int(len(tmp) / 2)) / len(tmp)
        nf = len(f)
        tmp1 = np.fft.fft(tmp)
        Four = abs(tmp1)
        return f, Four

#作图部分
def graph(e , n, u, name, method='a'):
    plt.rcParams['font.family'] = ['sans-serif']
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus']=False # 用来正常显示负号
    t=e.x

    plt.figure(figsize=(10,8))
    plt.subplot(3,1,1)
    plt.xticks([])
    if method == 'a':
        plt.plot(t,e.acc,label=e.id+" EW PGA="+str(e.get_PGA()))
    elif method == 'v':
        e.get_v()
        plt.plot(t,e.vel,label=e.id+" EW PGV="+str(e.get_PGV()))
    elif method == 'd':
        e.get_u()
        plt.plot(t,e.dis,label=e.id+" EW PGD="+str(e.get_PGD()))
    plt.legend()

    plt.subplot(3,1,2)
    plt.xticks([])
    if method == 'a':
        plt.plot(t, n.acc,label=n.id+" NS PGA="+str(n.get_PGA()))
    elif method == 'v':
        n.get_v()
        plt.plot(t, n.vel,label=e.id+" EW PGV="+str(n.get_PGV()))
    elif method == 'd':
        n.get_u()
        plt.plot(t, n.dis,label=e.id+" EW PGD="+str(n.get_PGD()))
    plt.legend()

    plt.subplot(3,1,3)
    plt.xlabel("t(s)")
    if method == 'a':
        plt.plot(t, u.acc,label=n.id+" UD PGA="+str(u.get_PGA()))
    elif method == 'v':
        u.get_v()
        plt.plot(t, u.vel,label=e.id+" UD PGV="+str(u.get_PGV()))
    elif method == 'd':
        u.get_u()
        plt.plot(t, u.dis,label=e.id+" UD PGD="+str(u.get_PGD()))
    plt.legend()

    plt.savefig(name)
    
def fyp(e,n,u, name, method = 'a', standard: str = None):
    plt.rcParams['font.family'] = ['sans-serif']
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus']=False # 用来正常显示负号
    plt.figure(figsize=(10,5))
    if method == 'a':
        if standard is not None:
            com=pd.read_csv(standard)    
            plt.plot(com.iloc[:,0],com.iloc[:,1],label="6度罕遇",linestyle="-")
            m=max(max(e.sa),max(n.sa),max(u.sa))/980
            print(m)
            print(com.iloc[:,2].max())
            if m >= com.iloc[:,2].max():
                plt.plot(com.iloc[:,0],com.iloc[:,2],label="7度罕遇",linestyle="-",marker="1")
            if m >= com.iloc[:,3].max():
                plt.plot(com.iloc[:,0],com.iloc[:,3],label="8度罕遇",marker="2")
            if m >= com.iloc[:,4].max():
                plt.plot(com.iloc[:,0],com.iloc[:,4],label="9度罕遇",marker="3")
        plt.plot(e.saT,e.sa/980,label=e.id+" EW",linestyle="solid",marker="^")
        plt.plot(e.saT,n.sa/980,label=n.id+" NS",linestyle='solid',marker="p")
        plt.plot(e.saT,u.sa/980,label=u.id+" UD",linestyle='dashdot',marker="o")
        plt.xlabel("T(s)")
        plt.ylabel("加速度反应谱(g)")
    elif method == 'v':
        plt.plot(e.saT,e.sv,label=e.id+" EW",linestyle="solid",marker="^")
        plt.plot(e.saT,n.sv,label=n.id+" NS",linestyle='solid',marker="p")
        plt.plot(e.saT,u.sv,label=u.id+" UD",linestyle='dashdot',marker="o")
        plt.xlabel("T(s)")
        plt.ylabel("速度反应谱(cm/s)")
    elif method == 'd':
        plt.plot(e.saT,e.sd,label=e.id+" EW",linestyle="solid",marker="^")
        plt.plot(e.saT,n.sd,label=n.id+" NS",linestyle='solid',marker="p")
        plt.plot(e.saT,u.sd,label=u.id+" UD",linestyle='dashdot',marker="o")
        plt.xlabel("T(s)")
        plt.ylabel("位移反应谱(cm)")
    plt.legend()
    plt.xlim(0,4)
    plt.savefig(name)
    
def fourht(e, n, u, name):
    plt.rcParams['font.family'] = ['sans-serif']
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus']=False # 用来正常显示负号
    plt.figure(figsize=(10,5))
    x,y=e.four()
    x1,y1=n.four()
    x2,y2=u.four()
    plt.plot(np.log10(x[1:]),np.log10(y[1:len(x)]),linewidth=0.5,label=e.id+" EW",linestyle="-")
    plt.plot(np.log10(x1[1:]),np.log10(y1[1:len(x1)]),linewidth=0.5,label=n.id+" NS",linestyle="--")
    plt.plot(np.log10(x2[1:]),np.log10(y2[1:len(x2)]),linewidth=0.5,label=u.id+" UD",linestyle=":")
    plt.xlim(np.log10(0.1),np.log10(50))
    plt.ylim(-4)
    x1=[np.log10(0.1),np.log10(1),np.log10(10),np.log10(50)]
    y1=['0.1','1','10','50']
    plt.xticks(x1,y1)
    x1=np.arange(-4,5)
    y1=["10e-4","10e-3","10e-2","10e-1","10e0","10e1","10e2","10e3","10e4"]
    plt.yticks(x1,y1)
    plt.legend()
    plt.xlabel("频率(Hz)")
    plt.ylabel("幅值谱(cm/s)")
    
    plt.savefig(name)
    return
    
class OmniSpectrum():

    def __init__(self, path: str = None, process: int = 4, period: int = 2, baseline_correction: bool = False, pool=None):  # 路径
        self.period = period
        self.process = process
        self.pool = pool
        self.baseline_correction = baseline_correction
        self.space_theta = np.radians(np.linspace(0, 360, 361))  # 转化为弧度
        self.space_r = np.linspace(0, self.period, self.period*10 + 1)
        self.r, self.theta = np.meshgrid(self.space_r, self.space_theta)
        self.sa = np.zeros(len(self.space_theta) * len(self.space_r)).reshape(len(self.space_theta), len(self.space_r))  # 数值矩阵 加速度
        self.sv = np.zeros(len(self.space_theta) * len(self.space_r)).reshape(len(self.space_theta), len(self.space_r))  # 速度
        self.su = np.zeros(len(self.space_theta) * len(self.space_r)).reshape(len(self.space_theta), len(self.space_r))  # 位移
        self.sa = multiprocessing.Manager().list(self.sa)
        self.sv = multiprocessing.Manager().list(self.sv)
        self.su = multiprocessing.Manager().list(self.su)
        self.path = path
        if path is not None: self.ew, self.ns, self.up, self.t = self.input_gmotion(path, mode = 1)
        
    def input_gmotion(self, path=None, ew=None, ns=None, up=None, mode: int = 1):  # 返回EW，NW,UP方向地震动数组和时间间隔t
        '''
        :param path:
        :param ew:
        :param ns:
        :param mode: 1:三向地震动同一文件，2；ew,ns地震动分别指定
        :return:
        '''
        if mode == 1:
            f = open(path, 'r')
            t = float(f.readline().split(',')[0])
            data = f.readlines()
            ew = []
            ns = []
            up = []
            for d in data:
                ew.append(float(d.split(',')[0]))
                ns.append(float(d.split(',')[1]))
                up.append(float(d.split(',')[2]))
            f.close()
        elif mode == 2:
            f = open(ew, 'r')
            t = float(f.readline().split(',')[0])
            data = f.readlines()
            ew = []
            for d in data:
                ew.append(float(d.split(',')[0]))
            f.close()
            f = open(ns, 'r')
            t = float(f.readline().split(',')[0])
            data = f.readlines()
            ew = []
            for d in data:
                ns.append(float(d.split(',')[0]))
            f.close()
        else:
            raise KeyError
        return ew, ns, up, t

    def get_single(self, direction):
        if direction == 'ew':
            return spectrum(self.ew, self.t, baseline_correction = self.baseline_correction)
        elif direction == 'ns':
            return spectrum(self.ns, self.t, baseline_correction = self.baseline_correction)
        elif direction == 'up':
            return spectrum(self.up, self.t, baseline_correction = self.baseline_correction)
        elif isinstance(direction, int):
            return spectrum(self.any_angle(self.ew, self.ns, np.radians(direction)), self.t, baseline_correction = self.baseline_correction)
        
    def any_angle(self, ew, ns, theta):  # EW,NS地震动,方向角，作用：返回任何方向的地震动
        return np.array(ew) * np.cos(theta) + np.array(ns) * np.sin(theta)

    def graph(self, theta, r, value, name):  # theta,r,value,命名（包含路径）
        fig, axes = plt.subplots(subplot_kw=dict(projection='polar'), figsize=(10, 10))
        contourplot = axes.contourf(theta, r, value, cmap=plt.cm.jet)
        plt.colorbar(contourplot, shrink=.6, pad=0.08)
        plt.savefig(name)

    def spect(self, direction):
        '''

        :param ew:
        :param ns:
        :param direction: 角度，0-360
        :return:
        '''
        # print(direction)
        new_a = self.any_angle(self.ew, self.ns, np.radians(direction))
        sp = self.get_single(direction)
        sp.get_sa(0, self.period + 0.1, 0.1)
        self.sa[direction] = sp.sa_truemax
        self.sa[direction + 180] = abs(sp.sa_truemin)
        self.sv[direction] = sp.sv_truemax
        self.sv[direction + 180] = abs(sp.sv_truemin)
        self.su[direction] = sp.sd_truemax
        self.su[direction + 180] = abs(sp.sd_truemin)
        # print('{}方向结束'.format(direction))
        return

    def any_angle_spectrum(self):
        if (not os.path.exists(self.path.split('.')[-2] + '_a.jpg')):
            for i in tqdm(range(0, 181)):
                self.spect(i)
#                 self.pool.apply_async(OmniSpectrum.spect, (new_a, t, self.period, i, self.value_a, self.value_v, self.value_u,), error_callback=errprint)
#             self.pool.close()
#             self.pool.join()
            self.graph(self.theta, self.r, self.sa, self.path.split('.')[-2] + '_a.jpg')
            self.graph(self.theta, self.r, self.sv, self.path.split('.')[-2] + '_v.jpg')
            self.graph(self.theta, self.r, self.su, self.path.split('.')[-2] + '_u.jpg')
            print("处理{}".format(self.path))
        else:
            print("跳过{}".format(self.path))
        return

class convert2csv():
    def get_data_china(self, path):
        f = open(path, 'r')
        d = f.readline()
        pos = d[:-1]
        for i in range(11):
            d = f.readline()
        t = float(d[53:59])
        for i in range(4):
            d = f.readline()
        data = f.readlines()
        sp = []
        for d in data:
            d1=d.split()
            for d11 in d1:
                sp.append(d11)
        f.close()
        return pos, t, sp
    
    def convert(self, path, output:str = None, mode:str = 'china'):
        if output is None: output = path
        if mode == 'china':
            data = {}
            for root,dirs,files in os.walk(path):
                for file in files:
                    if file.split('.')[-1] == 'dat':
                        p, t, sp = self.get_data_china(os.path.join(path,file))
                        data[p] = (t, sp)
            for d in data.keys():
                if d[-1] ==  '1':
                    f = open(os.path.join(output, '{}.csv'.format(d[0:6])), 'w')
                    f.write('{},{},\n'.format(data[d][0], d))
                    for i in range(len(data[d][1])):
                        ew = data[d][1]
                        tmp = d[0:-1] + '2'
                        ns = data[tmp][1]
                        tmp = d[0:-1] + '3'
                        up = data[tmp][1]
                        f.write('{},{},{}\n'.format(ew[i], ns[i],up[i]))
                    f.close()
        else:
            print('{}模式不存在'.format(mode))

def csv2docx(path, output, baseline_correction: bool = False, standard:str=os.path.join(os.path.dirname(__file__), 'standard.csv'), pool=None):
    doc=Document()
    for root,dirs,files in os.walk(path):
        for file in files:
            if file.split('.')[-1] == 'csv':
                print('{} start'.format(file))
                if (not os.path.exists(os.path.join(path, file).split('.')[-2] + '_acc.jpg')):
                    omni = OmniSpectrum(os.path.join(path, file), baseline_correction=baseline_correction)
                    e = omni.get_single('ew')
                    n = omni.get_single('ns')
                    u = omni.get_single('up')
                    e.get_u()
                    n.get_u()
                    u.get_u()
                    e.get_sa(0, 4, 0.1)
                    n.get_sa(0, 4, 0.1)
                    u.get_sa(0, 4, 0.1)
                    graph(e, n, u, os.path.join(path,file).split('.')[-2] + '_acc.jpg', method='a')
                    graph(e, n, u, os.path.join(path,file).split('.')[-2] + '_vel.jpg', method='v')
                    graph(e, n, u, os.path.join(path,file).split('.')[-2] + '_dis.jpg', method='d')
                    fyp(e, n, u, os.path.join(path,file).split('.')[-2] + '_sa.jpg', method='a', standard=standard)
                    fyp(e, n, u, os.path.join(path,file).split('.')[-2] + '_sv.jpg', method='v')
                    fyp(e, n, u, os.path.join(path,file).split('.')[-2] + '_sd.jpg', method='d')
                    fourht(e, n, u, os.path.join(path,file).split('.')[-2] + '_fourht.jpg')
                    omni.any_angle_spectrum()
                # word添加标题
                p = doc.add_paragraph(file.split('.')[0])
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tab =doc.add_table(rows=7,cols=2)

                miaoshu=["加速度时程曲线","速度时程曲线","位移时程曲线","加速度反应谱","速度反应谱","位移反应谱","傅里叶幅值谱","加速度全方位反应谱","速度全方位反应谱","位移全方位反应谱"]
                height=[7,7,6.38,5.2,5.2,5.2,5.2,6.68,6.68,6.68,6.68]
                pic_name = {
                    0 : os.path.join(path,file).split('.')[-2] + '_acc.jpg',
                    1 : os.path.join(path,file).split('.')[-2] + '_vel.jpg',
                    2 : os.path.join(path,file).split('.')[-2] + '_dis.jpg',
                    3 : os.path.join(path,file).split('.')[-2] + '_sa.jpg',
                    4 : os.path.join(path,file).split('.')[-2] + '_sv.jpg',
                    5 : os.path.join(path,file).split('.')[-2] + '_sd.jpg',
                    6 : os.path.join(path,file).split('.')[-2] + '_fourht.jpg',
                    7 : os.path.join(path,file).split('.')[-2] + '_a.jpg',
                    8 : os.path.join(path,file).split('.')[-2] + '_v.jpg',
                    9 : os.path.join(path,file).split('.')[-2] + '_u.jpg'
                }
                now_line=0
                for i in range(10):                    
                    if i not in [6,7,8,9]:
                        cell = tab.cell(i//2,i%2)
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(pic_name[i], height=shared.Cm(height[i]))
                        cell.add_paragraph(miaoshu[i])
                        cell.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #竖直居中
                        now_line=i//2
                    else:
                        now_line += 1
                        cell = tab.cell(now_line,0)
                        cell_new = cell.merge(tab.cell(now_line,1))
                        run = cell_new.paragraphs[0].add_run()
                        run.add_picture(pic_name[i], height=shared.Cm(height[i]))
                        cell_new.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell_new.add_paragraph(miaoshu[i])
                        cell_new.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #竖直居中
                        cell_new.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_page_break()
    doc.save(os.path.join(output, 'specturm.docx'))

if __name__ == '__main__':
    if len(sys.argv) > 1:
        parser = argparse.ArgumentParser(description='全方位反应谱生成程序')
        parser.add_argument('-i', '--input', type=str, help='输入文件路径', default=None)
        parser.add_argument('-o', '--output', type=str, help='输出文件路径', default=None)
        parser.add_argument('--baseline', action='store_true', help='基线修正', default=None)
        parser.add_argument('--thread', type=int, help='进程数', default=4)
        args = parser.parse_args()
    else:
        print('参数有误，--help查看帮助文档')

    assert args.input is not None, '输入路径为空！请使用--help指令查看帮助'
    path = args.input
    output = args.output if args.output is not None else args.input
    pool = multiprocessing.Pool(args.thread, maxtasksperchild=1)
    csv2docx(path, output, baseline_correction=args.baseline)